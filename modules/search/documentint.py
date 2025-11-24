# modules/search/documentint.py
import logging
import os
import re
from typing import List, Dict, Any
from pathlib import Path

# Importaciones condicionales para evitar errores al importar en entornos sin dependencias
try:
    from PyPDF2 import PdfReader

    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    from docx import Document

    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


def search_documents_in_pdf(pdf_path: str, search_terms: List[str]) -> List[Dict[str, Any]]:
    """
    Busca términos dentro de un documento PDF.
    Requiere: `pip install PyPDF2`
    """
    logger.info(f"Buscando términos '{search_terms}' en documento PDF: {pdf_path}")

    results = []

    if not PYPDF_AVAILABLE:
        logger.warning("PyPDF2 no está disponible. Instale con: pip install PyPDF2")
        return results

    try:
        if not os.path.exists(pdf_path):
            logger.warning(f"Documento PDF no encontrado: {pdf_path}")
            return results

        reader = PdfReader(pdf_path)
        total_pages = len(reader.pages)

        for page_num, page in enumerate(reader.pages):
            text_content = page.extract_text()

            if not text_content:  # Salta páginas sin texto
                continue

            for term in search_terms:
                # Encontrar todas las posiciones donde aparece el término (case insensitive)
                pattern = re.compile(re.escape(term), re.IGNORECASE)
                matches = [(m.start(), m.end()) for m in pattern.finditer(text_content)]

                if matches:
                    matched_positions = []
                    previews = []

                    # Para cada coincidencia, extraer fragmentos de contexto
                    for start, end in matches[:5]:  # Máximo 5 apariciones
                        # Contexto de 100 caracteres antes y después
                        context_start = max(0, start - 50)
                        context_end = min(len(text_content), end + 50)
                        preview = text_content[context_start:context_end].strip()

                        matched_positions.append(start)
                        previews.append({
                            "position": start,
                            "preview": preview,
                            "context_start": context_start,
                            "context_end": context_end
                        })

                    results.append({
                        "term": term,
                        "page": page_num + 1,
                        "total_matches": len(matches),
                        "positions": matched_positions[:5],
                        "previews": previews,
                        "context_length": len(text_content)
                    })

    except Exception as e:
        logger.exception(f"Error procesando PDF: {e}")
        return []

    return results


def search_docx_file(docx_path: str, search_term: str) -> List[Dict[str, Any]]:
    """
    Búsqueda de texto dentro de un archivo .docx (Word).
    Requiere: `pip install python-docx`
    """
    logger.info(f"Buscando '{search_term}' en documento Word: {docx_path}")

    results = []

    if not PYTHON_DOCX_AVAILABLE:
        logger.warning("python-docx no está disponible. Instale con: pip install python-docx")
        return results

    try:
        if not os.path.exists(docx_path):
            logger.warning(f"Documento DOCX no encontrado: {docx_path}")
            return results

        doc = Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        full_text_str = '\n'.join(full_text)

        if not full_text_str:  # No hay texto en el documento
            return results

        # Encontrar todas las posiciones del término buscado (case insensitive)
        pattern = re.compile(re.escape(search_term), re.IGNORECASE)
        matches = [(m.start(), m.end()) for m in pattern.finditer(full_text_str)]

        previews = []
        for i, (start, end) in enumerate(matches[:5]):
            # Generar vista previa con contexto
            context_start = max(0, start - 50)
            context_end = min(len(full_text_str), end + 50)
            preview = full_text_str[context_start:context_end].strip()

            previews.append({
                "position": start,
                "preview": preview,
                "context_start": context_start,
                "context_end": context_end
            })

        results = [{
            "term": search_term,
            "total_matches": len(matches),
            "previews": previews,
            "document_length": len(full_text_str)
        }]

    except Exception as e:
        logger.exception(f"Error procesando documento .docx: {e}")

    return results


def search_slideshare_document(document_id: str) -> Dict[str, Any]:
    """
    Busca un documento público en Slideshare (simulador).
    Devuelve metadata simulada de un documento público.
    """
    logger.info(f"Buscando documento en SlideShare con ID: {document_id}")

    # Simulación de datos de un documento público
    return {
        "document_id": document_id,
        "title": "Presentación Ejemplo sobre OSINT Básico",
        "author": "Ana Rodríguez",
        "description": "Una presentación de ejemplo sobre conceptos básicos de OSINT.",
        "tags": ["OSINT", "Cybersecurity", "Investigación"],
        "download_link": "https://slideshare.net/ejemplo/presentacion",
        "presentation_type": "PowerPoint",
        "slides": 24,
        "view_count": 1500,
        "publish_date": "2024-01-15",
        "categories": ["Tecnología", "Seguridad Informática"],
        "source": "slideshare"
    }


# Función auxiliar para obtener metadatos de documentos
def get_document_metadata(file_path: str) -> Dict[str, Any]:
    """
    Obtiene metadatos de un documento (PDF, DOCX, etc.)
    Requiere instalación correspondiente en función del formato.
    """
    logger.info(f"Extrayendo metadatos de: {file_path}")

    filename = os.path.basename(file_path)
    file_ext = os.path.splitext(filename)[1].lower()

    # Valores predeterminados
    metadata = {
        "filename": filename,
        "full_path": file_path,
        "extension": file_ext,
        "size_bytes": os.path.getsize(file_path) if os.path.exists(file_path) else 0,
        "created_time": "2024-08-10T14:30:00",
        "modified_time": "2024-10-25T09:00:00",
        "author": "Usuario Desconocido",
        "title": "",
        "subject": ""
    }

    try:
        if file_ext == '.pdf' and PYPDF_AVAILABLE:
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                info = reader.metadata
                if info:
                    # Normalizar nombres de campos (elimina '/' inicial)
                    clean_info = {
                        key.replace('/', ''): str(val)
                        for key, val in info.items()
                        if val is not None and val != ''
                    }
                    metadata.update(clean_info)

                    # Si no tiene fecha modificada, usar creación
                    if not metadata.get('ModifyDate') and metadata.get('CreateDate'):
                        metadata['modified_time'] = metadata.get('CreateDate')

        elif file_ext == '.docx' and PYTHON_DOCX_AVAILABLE:
            # Para MS Word (.docx)
            doc = Document(file_path)
            core_props = doc.core_properties

            if core_props:
                if core_props.author:
                    metadata['author'] = core_props.author

                if core_props.title:
                    metadata['title'] = core_props.title

                if core_props.subject:
                    metadata['subject'] = core_props.subject

                if core_props.created:
                    metadata['created_time'] = core_props.created.isoformat()

                if core_props.modified:
                    metadata['modified_time'] = core_props.modified.isoformat()

    except Exception as e:
        logger.exception(f"Error al extraer metadatos del documento: {e}")

    return metadata


def find_all_documents_in_directory(directory_path: str, extensions: List[str] = None) -> List[Dict[str, Any]]:
    """
    Busca todos los documentos de cierta extensión en un directorio.

    Args:
        directory_path: Ruta al directorio a escanear
        extensions: Lista de extensiones a buscar (ej: ['.pdf', '.docx'])

    Returns:
        Lista con la información de los documentos encontrados
    """
    if extensions is None:
        extensions = ['.pdf', '.docx', '.doc']

    found_docs = []
    directory = Path(directory_path)

    if not directory.exists():
        logger.warning(f"Directorio no encontrado: {directory_path}")
        return found_docs

    try:
        for ext in extensions:
            for file_path in directory.rglob(f"*{ext}"):
                if file_path.is_file():
                    metadata = get_document_metadata(str(file_path))
                    found_docs.append({
                        "path": str(file_path),
                        "size_mb": round(metadata["size_bytes"] / (1024 * 1024), 2),
                        "metadata": metadata,
                        "exists": True
                    })

    except Exception as e:
        logger.exception(f"Error al escanear directorio: {e}")

    return found_docs


def search_in_multiple_documents(doc_paths: List[str], search_terms: List[str]) -> Dict[str, Any]:
    """
    Busca términos en múltiples documentos

    Args:
        doc_paths: Lista de rutas a documentos
        search_terms: Lista de términos a buscar

    Returns:
        Diccionario con resultados agrupados por documento
    """
    all_results = {}

    for doc_path in doc_paths:
        filename = os.path.basename(doc_path)
        logger.info(f"Procesando documento: {filename}")

        results = []

        if doc_path.endswith('.pdf'):
            results = search_documents_in_pdf(doc_path, search_terms)
        elif doc_path.endswith(('.docx', '.doc')):
            # Se buscará el primer término en documentos .docx
            if search_terms:
                results = search_docx_file(doc_path, search_terms[0])

        all_results[filename] = {
            "path": doc_path,
            "results": results,
            "found_terms": len([r for r in results if r.get('total_matches', 0) > 0])
        }

    return all_results